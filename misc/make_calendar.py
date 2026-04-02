#!/usr/bin/env python3
"""Generate STUDY-CALENDAR.docx — landscape calendar of actual Semester 6 schedule."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colors (hex) ────────────────────────────────────────────────────────────
AC_BG     = "FFE4E4"
MT_BG     = "DCE8FF"  # deliberately using "E8" for slightly more blue
SEC_BG    = "FFF3DA"
BUF_BG    = "E1F5E1"
OFF_BG    = "F0F0F0"
WE_BG     = "E8E8E8"
HEAD_BG   = "333333"
EXAM_BG   = "F0E0F0"

AC_TEXT   = RGBColor(180, 40, 40)
MT_TEXT   = RGBColor(40, 100, 180)
SEC_TEXT  = RGBColor(140, 90, 20)
BUF_TEXT  = RGBColor(60, 130, 60)
EXAM_TEXT = RGBColor(190, 30, 30)
HOL_TEXT  = RGBColor(120, 80, 160)
GRAY_TEXT = RGBColor(130, 130, 130)
DAY_TEXT  = RGBColor(80, 80, 80)
WHITE     = RGBColor(255, 255, 255)
BLACK     = RGBColor(0, 0, 0)


def set_cell_bg(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_cell_text(cell, lines, date_color=DAY_TEXT, body_color=BLACK, bold_date=True, font_size=7):
    """Add text to a cell. First line = date (bold), rest = body."""
    cell.paragraphs[0].clear()
    for i, line in enumerate(lines):
        if i > 0:
            p = cell.add_paragraph()
        else:
            p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(font_size + 1.5)
        run = p.add_run(line)
        run.font.size = Pt(font_size if i > 0 else font_size + 0.5)
        run.font.color.rgb = date_color if i == 0 else body_color
        run.font.bold = bold_date if i == 0 else False
        run.font.name = "Calibri"


def add_header_cell(cell, text):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(7.5)
    run.font.color.rgb = WHITE
    run.font.bold = True
    run.font.name = "Calibri"
    set_cell_bg(cell, HEAD_BG)


def add_phase_cell(cell, lines):
    cell.paragraphs[0].clear()
    for i, line in enumerate(lines):
        if i > 0:
            p = cell.add_paragraph()
        else:
            p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(8)
        run = p.add_run(line)
        run.font.size = Pt(6.5 if i > 0 else 7.5)
        run.font.color.rgb = WHITE
        run.font.bold = True
        run.font.name = "Calibri"
    set_cell_bg(cell, HEAD_BG)


# ── Schedule data ───────────────────────────────────────────────────────────
# Each week: (phase_lines, [7 day tuples])
# day tuple: (bg_color, date_str, [lines], text_color, is_holiday)
schedule = [
    # ── Week 1: Mar 30 – Apr 5 ──
    (["AC", "Hackathon", "Wk 1"], [
        (OFF_BG,  "Mar 30", [""],                                        GRAY_TEXT, False),
        (OFF_BG,  "Mar 31", [""],                                        GRAY_TEXT, False),
        (OFF_BG,  "Apr 1",  ["Off"],                                     GRAY_TEXT, False),
        (AC_BG,   "Apr 2",  ["AC L1: Intro +", "Turing Machines"],       AC_TEXT,   False),
        (AC_BG,   "Apr 3",  ["AC L2: Church-", "Turing Thesis"],         AC_TEXT,   False),
        (AC_BG,   "Apr 4",  ["Work day"],                                GRAY_TEXT, False),
        (AC_BG,   "Apr 5",  ["AC L3:", "Computability"],                 AC_TEXT,   False),
    ]),
    # ── Week 2: Apr 6 – 12 ──
    (["AC", "Hackathon", "Wk 2"], [
        (AC_BG,   "Apr 6",  ["AC L4: Reductions"],                      AC_TEXT,   False),
        (AC_BG,   "Apr 7",  ["AC L5: P and NP"],                        AC_TEXT,   False),
        (OFF_BG,  "Apr 8",  ["Off"],                                     GRAY_TEXT, False),
        (AC_BG,   "Apr 9",  ["AC L6:", "NP-Completeness"],              AC_TEXT,   False),
        (AC_BG,   "Apr 10", ["AC L7: More", "NP-Completeness"],         AC_TEXT,   False),
        (WE_BG,   "Apr 11", ["Weekend"],                                 GRAY_TEXT, False),
        (WE_BG,   "Apr 12", ["Weekend"],                                 GRAY_TEXT, False),
    ]),
    # ── Week 3: Apr 13 – 19 ──
    (["AC", "Hackathon", "Wk 3"], [
        (AC_BG,   "Apr 13", ["AC L8: DPLL +", "Backtracking"],          AC_TEXT,   False),
        (AC_BG,   "Apr 14", ["AC L9: Advanced", "Graph Algorithms"],    AC_TEXT,   False),
        (OFF_BG,  "Apr 15", ["Off"],                                     GRAY_TEXT, False),
        (AC_BG,   "Apr 16", ["AC L10: Linear", "Programming"],          AC_TEXT,   False),
        (AC_BG,   "Apr 17", ["AC L11: Integer", "Linear Programming"],  AC_TEXT,   False),
        (WE_BG,   "Apr 18", ["Weekend"],                                 GRAY_TEXT, False),
        (WE_BG,   "Apr 19", ["Weekend"],                                 GRAY_TEXT, False),
    ]),
    # ── Week 4: Apr 20 – 26 ──
    (["AC", "Hackathon", "Wk 4"], [
        (AC_BG,   "Apr 20", ["AC L12: Amortized", "Analysis"],          AC_TEXT,   False),
        (AC_BG,   "Apr 21", ["AC Review:", "Exercises + weak", "topics from L1-L12"], AC_TEXT, False),
        (OFF_BG,  "Apr 22", ["Off"],                                     GRAY_TEXT, False),
        (AC_BG,   "Apr 23", ["L13 uploaded", "by professor"],           GRAY_TEXT, False),
        (AC_BG,   "Apr 24", ["AC L13: Dynamic", "Programming"],         AC_TEXT,   False),
        (WE_BG,   "Apr 25", ["Weekend"],                                 GRAY_TEXT, False),
        (WE_BG,   "Apr 26", ["Weekend"],                                 GRAY_TEXT, False),
    ]),
    # ── Week 5: Apr 27 – May 3 ──
    (["mtcps", "Hackathon", "Wk 1"], [
        (MT_BG,   "Apr 27", ["mtcps L1+L2:", "UPPAAL Tool + Intro"],    MT_TEXT,   False),
        (MT_BG,   "Apr 28", ["mtcps L3:", "Synchronous Model"],         MT_TEXT,   False),
        (OFF_BG,  "Apr 29", ["Off"],                                     GRAY_TEXT, False),
        (MT_BG,   "Apr 30", ["mtcps L4: Safety", "Requirements"],       MT_TEXT,   False),
        (MT_BG,   "May 1",  ["mtcps L5:", "Asynchronous Model"],        MT_TEXT,   False),
        (WE_BG,   "May 2",  ["Weekend"],                                 GRAY_TEXT, False),
        (WE_BG,   "May 3",  ["Weekend"],                                 GRAY_TEXT, False),
    ]),
    # ── Week 6: May 4 – 10 ──
    (["mtcps", "Hackathon", "Wk 2"], [
        (MT_BG,   "May 4",  ["mtcps L6:", "Timed Model"],               MT_TEXT,   False),
        (MT_BG,   "May 5",  ["mtcps L7:", "UPPAAL (deep)"],             MT_TEXT,   False),
        (OFF_BG,  "May 6",  ["Off"],                                     GRAY_TEXT, False),
        (MT_BG,   "May 7",  ["mtcps L8: Real-", "Time Scheduling"],     MT_TEXT,   False),
        (MT_BG,   "May 8",  ["mtcps L9: Buzzing", "Boys (Mini Proj.)"], MT_TEXT,   False),
        (WE_BG,   "May 9",  ["Weekend"],                                 GRAY_TEXT, False),
        (WE_BG,   "May 10", ["Weekend"],                                 GRAY_TEXT, False),
    ]),
    # ── Week 7: May 11 – 17 ──
    (["mtcps", "Hackathon", "Wk 3"], [
        (MT_BG,   "May 11", ["mtcps L10:", "Dynamical Systems"],         MT_TEXT,   False),
        (MT_BG,   "May 12", ["mtcps L11: Prev-", "Exam Walkthrough"],   MT_TEXT,   False),
        (OFF_BG,  "May 13", ["Off"],                                      GRAY_TEXT, False),
        (EXAM_BG, "May 14", ["PREV-EXAM 2023", "Timed: 4 hours"],       EXAM_TEXT, False),
        (EXAM_BG, "May 15", ["PREV-EXAM 2024", "Timed: 4 hours"],       EXAM_TEXT, False),
        (WE_BG,   "May 16", ["Weekend"],                                  GRAY_TEXT, False),
        (WE_BG,   "May 17", ["Weekend"],                                  GRAY_TEXT, False),
    ]),
    # ── Week 8: May 18 – 24 ──
    (["Security", "Sprint", "Wk 1"], [
        (SEC_BG,  "May 18", ["Sec L1: CIA Model", "+ L2: Access Ctrl"],  SEC_TEXT,  False),
        (SEC_BG,  "May 19", ["Sec L3: Secure Dev", "+ L4: Lang Sec"],    SEC_TEXT,  False),
        (OFF_BG,  "May 20", ["Off"],                                      GRAY_TEXT, False),
        (SEC_BG,  "May 21", ["Sec L5: Info Flow", "+ L6: Network Sec"],  SEC_TEXT,  False),
        (SEC_BG,  "May 22", ["Sec L7: Attack Tools", "+ L8: Rev. Eng."], SEC_TEXT,  False),
        (WE_BG,   "May 23", ["Weekend"],                                  GRAY_TEXT, False),
        (WE_BG,   "May 24", ["Weekend"],                                  GRAY_TEXT, False),
    ]),
    # ── Week 9: May 25 – 31 ──
    (["Sec Wk 2", "+", "Buffer"], [
        (SEC_BG,  "May 25", ["Sec L9: Hacking", "+ L10: Sec Analysis"], SEC_TEXT,  False),
        (SEC_BG,  "May 26", ["Sec L11: Counter-", "measures + L12: Rev"], SEC_TEXT, False),
        (OFF_BG,  "May 27", ["Off"],                                      GRAY_TEXT, False),
        (BUF_BG,  "May 28", ["Buffer: restudy", "earliest exam course"], BUF_TEXT,  False),
        (BUF_BG,  "May 29", ["Buffer: restudy", "same course"],          BUF_TEXT,  False),
        (WE_BG,   "May 30", ["Weekend"],                                  GRAY_TEXT, False),
        (WE_BG,   "May 31", ["Weekend"],                                  GRAY_TEXT, False),
    ]),
    # ── Week 10: Jun 1 – 7 ──
    (["Buffer", "+", "Exams"], [
        (BUF_BG,  "Jun 1",  ["Buffer: restudy", "2nd exam course"],      BUF_TEXT,  False),
        (BUF_BG,  "Jun 2",  ["Buffer: restudy", "same course"],          BUF_TEXT,  False),
        (OFF_BG,  "Jun 3",  ["Off"],                                      GRAY_TEXT, False),
        (BUF_BG,  "Jun 4",  ["Buffer: restudy", "3rd exam course"],      BUF_TEXT,  False),
        (BUF_BG,  "Jun 5",  ["Buffer: restudy", "same course"],          BUF_TEXT,  False),
        (WE_BG,   "Jun 6",  [""],                                         GRAY_TEXT, False),
        (WE_BG,   "Jun 7",  [""],                                         GRAY_TEXT, False),
    ]),
]


def main():
    doc = Document()

    # ── Page setup: landscape A4 ────────────────────────────────────────────
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.6)
    section.left_margin = Cm(0.8)
    section.right_margin = Cm(0.8)

    # ── Title ───────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("Semester 6 Study Calendar")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = "Calibri"
    run = title.add_run("    Apr 1 – Jun 5, 2026  ·  10 Weeks  ·  ")
    run.font.size = Pt(9)
    run.font.name = "Calibri"

    # Legend inline
    for label, color in [("AC (13 lec)", AC_TEXT), ("mtcps (11 lec)", MT_TEXT),
                          ("Security (12 lec)", SEC_TEXT), ("Buffer", BUF_TEXT)]:
        run = title.add_run(f"  ■ {label}")
        run.font.size = Pt(8)
        run.font.color.rgb = color
        run.font.bold = True
        run.font.name = "Calibri"

    run = title.add_run("   Hol.")
    run.font.size = Pt(7)
    run.font.color.rgb = HOL_TEXT
    run.font.bold = True
    run.font.name = "Calibri"
    run = title.add_run(" = holiday (study from home)")
    run.font.size = Pt(7)
    run.font.color.rgb = GRAY_TEXT
    run.font.name = "Calibri"

    # ── Table ───────────────────────────────────────────────────────────────
    cols = 8  # phase + 7 days
    rows = 1 + len(schedule)  # header + data
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Column widths
    phase_w = Cm(2.0)
    day_w = Cm(3.7)
    for row in table.rows:
        row.cells[0].width = phase_w
        for c in range(1, 8):
            row.cells[c].width = day_w

    # Set fixed row height for all rows
    for row_idx, row in enumerate(table.rows):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="680" w:hRule="atLeast"/>')
        trPr.append(trHeight)

    # ── Header row ──────────────────────────────────────────────────────────
    headers = ["Phase", "Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
    for i, h in enumerate(headers):
        add_header_cell(table.rows[0].cells[i], h)

    # ── Data rows ───────────────────────────────────────────────────────────
    for week_idx, (phase_lines, days) in enumerate(schedule):
        row = table.rows[week_idx + 1]

        # Phase cell
        add_phase_cell(row.cells[0], phase_lines)

        # Day cells
        for day_idx, day_data in enumerate(days):
            cell = row.cells[day_idx + 1]
            bg, date_str, lines, text_color, is_holiday = day_data

            if bg is None:
                # Empty/unused cell
                set_cell_bg(cell, WE_BG)
                continue

            set_cell_bg(cell, bg)

            if not date_str:
                continue

            # Build cell content
            all_lines = []
            date_label = date_str
            if is_holiday:
                date_label += "  Hol."
            all_lines.append(date_label)
            if lines:
                all_lines.extend([l for l in lines if l])

            date_c = DAY_TEXT
            body_c = text_color if text_color else BLACK

            cell.paragraphs[0].clear()
            for i, line in enumerate(all_lines):
                if i > 0:
                    p = cell.add_paragraph()
                else:
                    p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = Pt(8.5)

                if i == 0:
                    # Date line — possibly with "Hol." suffix
                    if is_holiday:
                        run = p.add_run(date_str + "  ")
                        run.font.size = Pt(7.5)
                        run.font.color.rgb = date_c
                        run.font.bold = True
                        run.font.name = "Calibri"
                        run2 = p.add_run("Hol.")
                        run2.font.size = Pt(6)
                        run2.font.color.rgb = HOL_TEXT
                        run2.font.bold = True
                        run2.font.name = "Calibri"
                    else:
                        run = p.add_run(date_str)
                        run.font.size = Pt(7.5)
                        run.font.color.rgb = date_c
                        run.font.bold = True
                        run.font.name = "Calibri"
                else:
                    run = p.add_run(line)
                    run.font.size = Pt(7)
                    run.font.color.rgb = body_c
                    run.font.name = "Calibri"
                    # Bold lecture identifiers
                    if any(line.startswith(prefix) for prefix in
                           ["AC L", "mtcps L", "Sec L", "PREV-EXAM", "Buffer"]):
                        run.font.bold = True

    # ── Summary line ────────────────────────────────────────────────────────
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.paragraph_format.space_before = Pt(4)
    summary.paragraph_format.space_after = Pt(0)

    parts = [
        ("AC Hackathon: Apr 2–24 (13 lec, 4 wk)  |  1 lec/day  |  A4 at exam", AC_TEXT),
        ("      ", BLACK),
        ("mtcps Hackathon: Apr 27–May 15 (11 lec + 2 prev-exams)  |  Open-book", MT_TEXT),
        ("      ", BLACK),
        ("Security: May 18–26 (2 lec/day)", SEC_TEXT),
        ("      ", BLACK),
        ("Buffer: May 28 → Exams", BUF_TEXT),
    ]
    for text, color in parts:
        run = summary.add_run(text)
        run.font.size = Pt(7)
        run.font.color.rgb = color
        run.font.bold = True
        run.font.name = "Calibri"

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(1)
    run = footer.add_run(
        "Pattern: Mon–Tue–Thu–Fri study  ·  Wed always off  ·  "
        "Buffer = 2 days/course before each exam  ·  Jawad M.K. Qayyum"
    )
    run.font.size = Pt(6.5)
    run.font.color.rgb = GRAY_TEXT
    run.font.name = "Calibri"

    # ── Save ────────────────────────────────────────────────────────────────
    out_path = os.path.join(os.path.dirname(__file__), "STUDY-CALENDAR.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
